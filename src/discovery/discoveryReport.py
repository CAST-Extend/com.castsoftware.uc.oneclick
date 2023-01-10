from discovery.sourceValidation import SourceValidation 
from pandas import DataFrame,read_excel
from os.path import abspath
import docx
from config import Config
from util import convert_LOC

#todo: add totals to the cloc report (d1)
#todo: if SQL problems, add bullet under SQL Delivery to describe (d1)

class DiscoveryReport(SourceValidation):

    def __init__(cls, config:Config, log_level:int):
        super().__init__(config,cls.__class__.__name__,log_level)


    def cloc_report(cls,file:str,sheet:str):
        
        # read by 'Stats Before Code CleanUP' sheet of an Cloc_Output excel file
        df = read_excel(file,sheet_name=sheet)
        df = df[df['LANGUAGE']!='SUM:'] # remove the total row

        # #extract total cout of applicable code from 'Stats Before Code CleanUP' sheet of an Cloc_Output excel file
        # filt=(df['APPLICABLE']==True)
        # pre_LOC=(df.loc[filt,'CODE'].sum())
        # pre_LOC=pre_LOC//1000

        return df


    def run(cls,config:Config):

        cloc_report = abspath(f'{config.report}/{config.project_name}/{config.project_name}-cloc.xlsx')
        discovery_report = abspath(f'{config.report}/{config.project_name}/{config.project_name}-source-code-discovery.docx')
        # create an instance of a word document
        doc = docx.Document()
        # add a heading of level 0 (largest heading)
        doc.add_heading('Source code discovery report', 0)
        doc_para = doc.add_paragraph(f'Please find below the completed discovery for the Project {config.project_name}')

        for appl in config.application:
            cls._log.info(f'Running {cls.__class__.__name__} for {appl}')

            sql_report = abspath(f'{config.report}/{config.project_name}/{appl}-SQLReport.xlsx')

            # read by 'Stats Before Code CleanUP' sheet of an Cloc_Output excel file
            before_df = cls.cloc_report(cloc_report,f'Before-Cleanup({appl})')

            filt=(before_df['APPLICABLE']==False)
            l=before_df.loc[filt,['LANGUAGE','CODE']].sort_values(by=['CODE'], ascending=False) 
            non_code=''
            if len(l) == 1:
                non_code= str(l.iloc[0]['CODE']//1000) +' KLOC of '+ l.iloc[0]['LANGUAGE']
            elif len(l) > 1:
                non_code= str(l.iloc[0]['CODE']//1000) +' KLOC of '+ l.iloc[0]['LANGUAGE'] +' non code files and ~'+ str(l.iloc[1]['CODE']//1000) +' KLOC of '+ l.iloc[1]['LANGUAGE']

            # read by 'Stats After Code CleanUP' sheet of an Cloc_Output excel file
            after_df = cls.cloc_report(cloc_report,f'After-Cleanup({appl})')

            # read by 'Summary' sheet of an SQL_Output excel file
            sql_df = read_excel(sql_report,sheet_name='Summary')
            sql_df = sql_df[sql_df['Total']>0]

            doc.add_heading(f'Application {appl} :', 1)

            # out of scope code base
            total = int(before_df['CODE'].sum())
            total, unit = convert_LOC(total)
            doc.add_paragraph(f"This delivery contains a total of {len(before_df)-1} languages with a total of {total} {unit}.",style='List Bullet')

            bsuport = before_df[before_df['APPLICABLE']==True]
            total = int(bsuport['CODE'].sum())
            total, unit = convert_LOC(total)

            lang_list = list(bsuport['LANGUAGE'])
            lang_list[-1]=f'and {lang_list[-1]}'
            bsup_lang = ', '.join(lang_list)
            doc.add_paragraph(f"{len(bsuport)} are supported by CAST, {bsup_lang} containing {total} {unit} and are in scope.",style='List Bullet 2')

            nsuport = before_df[before_df['APPLICABLE']==False]
            total = int(bsuport['CODE'].sum())
            total, unit = convert_LOC(total)

            lang_list = list(nsuport['LANGUAGE'])
            lang_list[-1]=f'and {lang_list[-1]}'
            bsup_lang = ', '.join(lang_list)
            doc.add_paragraph(f"The remaining {len(nsuport)} are unsupported, {bsup_lang} containing ({total} {unit}) and considered out of scope.",style='List Bullet 2')

            # in scope code base
            asuport = after_df[before_df['APPLICABLE']==True]
            total = int(asuport['CODE'].sum())
            total, unit = convert_LOC(total)
            doc.add_paragraph(f'After removing all Sample, Test and other non production related code there is ({total} {unit}) in scope for this project',style='List Bullet 2')

            line = []
            for index, row in sql_df.iterrows():
                line.append(f"{row['Unique']} {row['Catagory']}")
                if row['Catagory'] == 'SQL files':
                    line[-1]=f'in {line[-1]}'
            doc.add_paragraph(f"The delivery also inclues, {', '.join(line)}.",style='List Bullet')

            doc.add_paragraph("Here is a high-level LOC per core techno of the in-scope code taken using CLOC utility",style='List Bullet')

            after_df = after_df[['LANGUAGE','FILES','CODE']]
            
            t = doc.add_table(after_df.shape[0]+1, after_df.shape[1])
            # add the header rows.
            for j in range(after_df.shape[-1]):
                t.cell(0,j).text = after_df.columns[j]

            # add the rest of the data frame
            for i in range(after_df.shape[0]):
                for j in range(after_df.shape[-1]):
                    number = after_df.values[i,j]
                    if type(number) is int:
                        number = "{:,}".format(number)
                    t.cell(i+1,j).text = number
            
            # Adding style to a table
            t.style = 'Light List Accent 1'

        # now save the document to a location
        doc.save(discovery_report)
        cls._log.info(f'Source Code Discovery report has been written to {discovery_report}')
    pass
