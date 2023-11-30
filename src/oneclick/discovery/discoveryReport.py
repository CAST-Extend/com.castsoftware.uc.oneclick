import pandas
from oneclick.discovery.sourceValidation import SourceValidation 
from pandas import DataFrame,read_excel
from os.path import abspath,exists
from oneclick.config import Config
from cast_common.util import convert_LOC,list_to_text

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.shared import Length
from tqdm import tqdm

#todo: add totals to the cloc report (d1)
#todo: if SQL problems, add bullet under SQL Delivery to describe (d1)

class DiscoveryReport(SourceValidation):

    def __init__(cls, config:Config, log_level:int):
        super().__init__(config,cls.__class__.__name__,log_level)


    def cloc_report(cls,file:str,sheet:str):
        
        # read by 'Stats Before Code CleanUP' sheet of an Cloc_Output excel file
        df = DataFrame()
        try:
            if exists(file):
                df = read_excel(file,sheet_name=sheet)
        except ValueError as ex:
            cls._log.warning(f'No cloc informaiton found for {sheet}')

        return df


    def run(cls,config:Config):

        cloc_report = abspath(f'{config.report}/{config.project_name}/{config.project_name}-cloc.xlsx')
        discovery_report = abspath(f'{config.report}/{config.project_name}/{config.project_name}-source-code-discovery.docx')
        # create an instance of a word document
        doc = Document()
        # add a heading of level 0 (largest heading)
        doc.add_heading('Source code discovery report', 0)
        #doc_para = doc.add_paragraph(f'Please find below the completed discovery for the Project {config.project_name}')
        doc_para = doc.add_paragraph(f'Please find below the completed discovery for Project {config.project_name}. Note that we used CLOC for quick discovery and completeness verification. The actual lines of code discovered by the CAST solution may differ slightly from below. ')
        doc.add_heading('Questions', 1)
        doc.add_paragraph('<Specific questions to be added manually if needed. Remove this section if no questions>')

        doc.add_heading('Observation by Application', 1)

        for appl in tqdm(config.application,desc='Discovery',leave=True, position=1):
            cls._log.info(f'Running {cls.__class__.__name__} for {appl}')

            sql_report = abspath(f'{config.report}/{config.project_name}/{appl}/{appl}-SQLReport.xlsx')

            doc.add_heading(f'Application {appl}:', 2)

            before_df = cls.cloc_report(cloc_report,f'{appl}-Before')
            if before_df.empty:
                doc.add_paragraph(f'This application contains NO source code')
            else:
                # read by 'Stats Before Code CleanUP' sheet of an Cloc_Output excel file
                before_df=before_df[before_df['LANGUAGE'] != 'Totals']
                before_df[before_df.columns]=before_df.apply(lambda x: x.str.strip() if isinstance(x, str) else x)
                #print(before_df)

                filt=(before_df['APPLICABLE']==False)
                l=before_df.loc[filt,['LANGUAGE','CODE']].sort_values(by=['CODE'], ascending=False) 
                non_code=''
                if len(l) == 1:
                    str(l.iloc[0]['CODE']//1000) +' KLOC of '+ l.iloc[0]['LANGUAGE']
                elif len(l) > 1:
                    non_code= str(l.iloc[0]['CODE']//1000) +' KLOC of '+ l.iloc[0]['LANGUAGE'].strip() +' non code files and ~'+ str(l.iloc[1]['CODE']//1000) +' KLOC of '+ l.iloc[1]['LANGUAGE']

                # read by 'Stats After Code CleanUP' sheet of an Cloc_Output excel file
                after_df = cls.cloc_report(cloc_report,f'{appl}-After')
                #print(after_df)

                # read by 'Summary' sheet of an SQL_Output excel file
                if exists(sql_report):
                    sql_df = read_excel(sql_report,sheet_name='Summary')
                else:
                    sql_df = DataFrame(columns=['Name','Total','Unique','Dups'])
                    cls._log.warning('No SQL found')
                #sql_df = sql_df[sql_df['Total']>0]

                # out of scope code base
                langs=len(before_df)
                total, unit = convert_LOC(int(before_df['CODE'].sum()))
                doc.add_paragraph(f"This delivery contains {langs} discovered technolog{'ies' if langs > 1 else 'y'}/extension{'s' if langs > 1 else ''}  and a total of {total} {unit}.",style='List Bullet')

                bsuport = after_df[after_df['APPLICABLE']==True]
                total = int(bsuport['CODE'].sum())
                total, unit = convert_LOC(total)

                lang_list = list(bsuport['LANGUAGE'].str.strip())
                langs = len(lang_list)
                if langs >1:
                    lang_list[-1]=f'and {lang_list[-1]}'
                bsup_lang = ', '.join(lang_list)
                doc.add_paragraph(f"{len(bsuport)} technolog{'ies' if langs > 1 else 'y'}/extension{'s are' if langs > 1 else ' is'} relevant for the CAST analysis, which include {bsup_lang}.",style='List Bullet')

                nsuport = after_df[after_df['APPLICABLE']==False]
                total = int(bsuport['CODE'].sum())
                total, unit = convert_LOC(total)

                lang_list = list(nsuport['LANGUAGE'].str.strip())
                langs = len(lang_list)
                if langs > 1:
                    lang_list[-1]=f'and {lang_list[-1]}'
                bsup_lang = ', '.join(lang_list)
                doc.add_paragraph(f"The remaining technolog{'ies' if langs > 1 else 'y'}/extension{'s are' if langs > 1 else ' is'} not relevant and will not be analyzed. These include {bsup_lang}.",style='List Bullet')

                # in scope code base
                # asuport = after_df[before_df['APPLICABLE']==True]
                # total = int(asuport['CODE'].sum())
                # total, unit = convert_LOC(total)
                #doc.add_paragraph(f'After removing all Sample, Test and other non production related code there is ({total} {unit}) in scope for this project',style='List Bullet 2')
                #print(sql_df)
                create_df=sql_df[sql_df['Name'].str.startswith('Create')]
                if len(create_df):
                    sql_items = []
                    create_df.sort_values(['Unique'],ascending=False,inplace=True)
                    for index,item in create_df.iterrows():
                        count = int(item['Unique'])             # get the unique count from dataframe
                        name = item['Name'].split(" ")[1]       # remove the word Create
                        if count == 1: name = name[:-1]         # only one item, remove make name singular
                        sql_items.append(f'{count} {name}')
                    doc.add_paragraph(f'The database contains a total of {list_to_text(sql_items)}.',style='List Bullet')
                else:
                    doc.add_paragraph(f'No database objects found.',style='List Bullet')

                removed_code=before_df['CODE'].sum()-after_df['CODE'].sum()
                doc.add_paragraph(f'We removed {round(removed_code/1000,1)} KLOC of sample, test and other non-production code from the source code.',style='List Bullet')
                after_df=after_df.drop(after_df[after_df['APPLICABLE']==0.0].index)
                total = after_df['CODE'].sum()//1000
                doc.add_paragraph(f"Here is a high-level summary of the source code that will be analyzed by CAST, totaling {total} KLOC.",style='List Bullet')

                after_df = after_df[['LANGUAGE','FILES','CODE']]
                after_df = after_df[:-1]
                files_count=after_df['FILES'].sum()
                code_count=after_df['CODE'].sum()
                new_row = pandas.DataFrame({'LANGUAGE':'Totals', 'FILES':[files_count], 'CODE':[code_count]})
                after_df = pandas.concat([after_df,new_row])
                
                t = doc.add_table(after_df.shape[0]+1, after_df.shape[1])
                # add the header rows.
                for j in range(after_df.shape[-1]):
                    t.cell(0,j).text = after_df.columns[j].strip()
                    if j > 0:
                        t.cell(0,j).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

                # add the rest of the data frame
                for i in range(after_df.shape[0]):
                    for j in range(after_df.shape[-1]):
                        number = after_df.values[i,j]
                        if type(number) is int:
                            number = "{:,}".format(number)
                        t.cell(i+1,j).text = number
                        if j > 0:
                            t.cell(i+1,j).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
                make_rows_bold(t.rows[len(after_df)])

                # Adding style to a table
                t.style = 'Medium Shading 1 Accent 1'

        # now save the document to a location
        doc.save(discovery_report)
        cls._log.info(f'Source Code Discovery report has been written to {discovery_report}')


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
